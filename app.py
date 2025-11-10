from flask import Flask, render_template, request, redirect, url_for, session, abort, send_file
import sqlite3
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, date
from functools import wraps
import io
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuration ---
app = Flask(__name__)
app.secret_key = 'your_super_secret_club_key_12345'
USERS_DB = 'users.db'
LOGS_DB = 'logs.db'


# --- Database Connection Helpers ---
def get_users_db():
    conn = sqlite3.connect(USERS_DB)
    conn.row_factory = sqlite3.Row
    return conn


def get_logs_db():
    conn = sqlite3.connect(LOGS_DB)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with app.app_context():
        # Initialize USERS_DB
        user_db = get_users_db()
        user_db.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                student_name TEXT NOT NULL,
                is_admin INTEGER NOT NULL DEFAULT 0 
            )
        ''')

        try:
            user_db.execute('SELECT is_admin FROM users LIMIT 1')
        except sqlite3.OperationalError:
            user_db.execute('ALTER TABLE users ADD COLUMN is_admin INTEGER NOT NULL DEFAULT 0')

        # Insert admin user
        ADMIN_USERNAME = "admin"
        ADMIN_NAME = "Club Administrator"
        ADMIN_PASSWORD = "adminpassword"
        hashed_password = generate_password_hash(ADMIN_PASSWORD)

        try:
            user_db.execute(
                'INSERT INTO users (username, password_hash, student_name, is_admin) VALUES (?, ?, ?, ?)',
                (ADMIN_USERNAME, hashed_password, ADMIN_NAME, 1)
            )
            print(f"✅ Admin user '{ADMIN_USERNAME}' added.")
        except sqlite3.IntegrityError:
            print(f"⚠️ Admin user '{ADMIN_USERNAME}' already exists.")
        user_db.commit()

        # Initialize LOGS_DB with enhanced schema
        log_db = get_logs_db()
        log_db.execute('''
            CREATE TABLE IF NOT EXISTS logs (
                id INTEGER PRIMARY KEY,
                user_id INTEGER NOT NULL,
                date TEXT NOT NULL,
                work_done TEXT NOT NULL,
                check_in_time TEXT,
                check_out_time TEXT,
                timestamp TEXT,
                is_archived INTEGER DEFAULT 0
            )
        ''')

        # Add new columns if they don't exist (migration)
        try:
            log_db.execute('SELECT check_in_time FROM logs LIMIT 1')
        except sqlite3.OperationalError:
            log_db.execute('ALTER TABLE logs ADD COLUMN check_in_time TEXT')
            log_db.execute('ALTER TABLE logs ADD COLUMN check_out_time TEXT')
            log_db.execute('ALTER TABLE logs ADD COLUMN is_archived INTEGER DEFAULT 0')

        log_db.commit()


# --- Security Decorator ---
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        if session.get('is_admin', 0) != 1:
            return abort(403)
        return f(*args, **kwargs)

    return decorated_function


# --- Archive Old Logs Function ---
def archive_old_logs():
    """Archives logs from previous days"""
    today = date.today().strftime('%Y-%m-%d')
    log_db = get_logs_db()
    log_db.execute('UPDATE logs SET is_archived = 1 WHERE date < ? AND is_archived = 0', (today,))
    log_db.commit()


# --- Flask Routes ---
@app.route('/')
def index():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    if session.get('is_admin', 0) == 1:
        return redirect(url_for('dashboard'))
    else:
        return redirect(url_for('submit_log'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        db = get_users_db()
        user = db.execute(
            'SELECT id, username, password_hash, student_name, is_admin FROM users WHERE username = ?', (username,)
        ).fetchone()

        if user and check_password_hash(user['password_hash'], password):
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['student_name'] = user['student_name']
            session['is_admin'] = user['is_admin']

            if user['is_admin'] == 1:
                return redirect(url_for('dashboard'))
            else:
                return redirect(url_for('submit_log'))
        else:
            error = 'Invalid Username or Password'

    return render_template('login.html', error=error)


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


@app.route('/dashboard', methods=['GET'])
@admin_required
def dashboard():
    archive_old_logs()

    student_name = session['student_name']
    user_db = get_users_db()
    log_db = get_logs_db()

    # Get filter parameters
    view_mode = request.args.get('view', 'today')
    selected_date = request.args.get('date', date.today().strftime('%Y-%m-%d'))

    # Fetch logs based on filter
    if view_mode == 'today':
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs WHERE date = ? AND is_archived = 0 ORDER BY timestamp DESC',
            (date.today().strftime('%Y-%m-%d'),)
        ).fetchall()
    elif view_mode == 'date':
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs WHERE date = ? ORDER BY timestamp DESC',
            (selected_date,)
        ).fetchall()
    elif view_mode == 'archived':
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs WHERE is_archived = 1 ORDER BY date DESC, timestamp DESC'
        ).fetchall()
    else:  # all
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs ORDER BY date DESC, timestamp DESC'
        ).fetchall()

    users_data = user_db.execute('SELECT id, student_name FROM users').fetchall()
    user_map = {user['id']: user['student_name'] for user in users_data}

    combined_logs = []
    for log in logs_data:
        combined_logs.append({
            'student_name': user_map.get(log['user_id'], 'Unknown User'),
            'date': log['date'],
            'work_done': log['work_done'],
            'check_in_time': log['check_in_time'],
            'check_out_time': log['check_out_time'],
            'timestamp': log['timestamp']
        })

    # Fetch all members for the member list
    all_members = user_db.execute(
        'SELECT id, username, student_name, is_admin FROM users ORDER BY is_admin DESC, student_name ASC'
    ).fetchall()

    # Count total members (excluding admins for accurate member count)
    total_members = user_db.execute('SELECT COUNT(*) as count FROM users WHERE is_admin = 0').fetchone()['count']
    all_users = user_db.execute('SELECT COUNT(*) as count FROM users').fetchone()['count']

    is_admin = session.get('is_admin', 0)
    return render_template('dashboard.html', logs=combined_logs, student_name=student_name,
                           is_admin=is_admin, view_mode=view_mode, selected_date=selected_date,
                           members=all_members, total_members=total_members, all_users=all_users)


@app.route('/submit_log', methods=['GET', 'POST'])
def submit_log():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    if session.get('is_admin', 0) == 1:
        return redirect(url_for('dashboard'))

    user_id = session['user_id']
    student_name = session['student_name']
    log_db = get_logs_db()

    # Check for today's log
    today_str = date.today().strftime('%Y-%m-%d')
    existing_log = log_db.execute(
        'SELECT id, check_in_time, check_out_time FROM logs WHERE user_id = ? AND date = ?',
        (user_id, today_str)
    ).fetchone()

    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'check_in' and not existing_log:
            check_in_time = datetime.now().strftime('%H:%M:%S')
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            log_db.execute(
                'INSERT INTO logs (user_id, date, work_done, check_in_time, timestamp, is_archived) VALUES (?, ?, ?, ?, ?, 0)',
                (user_id, today_str, '', check_in_time, timestamp)
            )
            log_db.commit()
            return redirect(url_for('submit_log'))

        elif action == 'check_out' and existing_log:
            check_out_time = datetime.now().strftime('%H:%M:%S')
            log_db.execute(
                'UPDATE logs SET check_out_time = ? WHERE id = ?',
                (check_out_time, existing_log['id'])
            )
            log_db.commit()
            return redirect(url_for('submit_log'))

        elif action == 'submit_work' and existing_log:
            work_done = request.form['work_done']
            log_db.execute(
                'UPDATE logs SET work_done = ? WHERE id = ?',
                (work_done, existing_log['id'])
            )
            log_db.commit()
            return render_template('submit_log.html', success=True, student_name=student_name,
                                   existing_log=existing_log)

    return render_template('submit_log.html', student_name=student_name, existing_log=existing_log)


@app.route('/admin/add_member', methods=['GET', 'POST'])
@admin_required
def admin_add_member():
    error = None
    success_message = None

    if request.method == 'POST':
        new_username = request.form['username']
        new_name = request.form['student_name']
        temp_password = request.form['temp_password']

        if not new_username or not new_name or not temp_password:
            error = "All fields are required."
        else:
            db = get_users_db()
            hashed_password = generate_password_hash(temp_password)

            try:
                db.execute(
                    'INSERT INTO users (username, password_hash, student_name, is_admin) VALUES (?, ?, ?, 0)',
                    (new_username, hashed_password, new_name)
                )
                db.commit()
                success_message = f"✅ Member **{new_name}** added successfully! Username: **{new_username}**"
            except sqlite3.IntegrityError:
                error = f"User with username '{new_username}' already exists."
            except Exception as e:
                error = f"Database error: {e}"

    # Fetch all members for display
    db = get_users_db()
    all_members = db.execute(
        'SELECT id, username, student_name, is_admin FROM users ORDER BY is_admin DESC, student_name ASC'
    ).fetchall()

    return render_template('admin_member.html', error=error, success_message=success_message, members=all_members)


@app.route('/admin/edit_member/<int:member_id>', methods=['GET', 'POST'])
@admin_required
def edit_member(member_id):
    db = get_users_db()
    error = None
    success_message = None

    if request.method == 'POST':
        new_password = request.form.get('new_password')

        if new_password:
            hashed_password = generate_password_hash(new_password)
            try:
                db.execute(
                    'UPDATE users SET password_hash = ? WHERE id = ?',
                    (hashed_password, member_id)
                )
                db.commit()
                success_message = "✅ Password updated successfully!"
            except Exception as e:
                error = f"Error updating password: {e}"
        else:
            error = "Password cannot be empty."

    member = db.execute('SELECT id, username, student_name, is_admin FROM users WHERE id = ?', (member_id,)).fetchone()

    if not member:
        abort(404)

    return render_template('edit_member.html', member=member, error=error, success_message=success_message)


@app.route('/admin/delete_member/<int:member_id>', methods=['POST'])
@admin_required
def delete_member(member_id):
    user_db = get_users_db()
    log_db = get_logs_db()

    # Prevent deleting yourself or the main admin
    if member_id == session.get('user_id'):
        return redirect(url_for('admin_add_member'))

    # Check if user is admin
    member = user_db.execute('SELECT is_admin, username, student_name FROM users WHERE id = ?', (member_id,)).fetchone()
    if member and member['is_admin'] == 1:
        return redirect(url_for('admin_add_member'))

    try:
        # Delete all logs associated with this member
        log_db.execute('DELETE FROM logs WHERE user_id = ?', (member_id,))
        log_db.commit()

        # Delete the user account (username becomes available again)
        user_db.execute('DELETE FROM users WHERE id = ?', (member_id,))
        user_db.commit()

        print(
            f"✅ Successfully deleted member: {member['student_name']} (username: {member['username']} is now available)")
    except Exception as e:
        print(f"Error deleting member: {e}")

    return redirect(url_for('admin_add_member'))


@app.route('/export/excel')
@admin_required
def export_excel():
    view_mode = request.args.get('view', 'today')
    selected_date = request.args.get('date', date.today().strftime('%Y-%m-%d'))

    log_db = get_logs_db()
    user_db = get_users_db()

    # Fetch logs based on view mode
    if view_mode == 'today':
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs WHERE date = ? AND is_archived = 0',
            (date.today().strftime('%Y-%m-%d'),)
        ).fetchall()
        filename = f"club_logs_today_{date.today()}.xlsx"
    elif view_mode == 'date':
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs WHERE date = ?',
            (selected_date,)
        ).fetchall()
        filename = f"club_logs_{selected_date}.xlsx"
    else:
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs ORDER BY date DESC'
        ).fetchall()
        filename = "club_logs_all.xlsx"

    users_data = user_db.execute('SELECT id, student_name FROM users').fetchall()
    user_map = {user['id']: user['student_name'] for user in users_data}

    # Create Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Club Logs"

    # Headers
    headers = ['Student Name', 'Date', 'Check-In Time', 'Check-Out Time', 'Work Summary', 'Submission Time']
    ws.append(headers)

    # Style headers
    header_fill = PatternFill(start_color="495057", end_color="495057", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # Add data
    for log in logs_data:
        ws.append([
            user_map.get(log['user_id'], 'Unknown'),
            log['date'],
            log['check_in_time'] or 'N/A',
            log['check_out_time'] or 'N/A',
            log['work_done'] or 'No work summary',
            log['timestamp']
        ])

    # Adjust column widths
    for column in ws.columns:
        max_length = 0
        column = list(column)
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column[0].column_letter].width = adjusted_width

    # Save to bytes
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(output, download_name=filename, as_attachment=True,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


@app.route('/export/word')
@admin_required
def export_word():
    view_mode = request.args.get('view', 'today')
    selected_date = request.args.get('date', date.today().strftime('%Y-%m-%d'))

    log_db = get_logs_db()
    user_db = get_users_db()

    # Fetch logs
    if view_mode == 'today':
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs WHERE date = ? AND is_archived = 0',
            (date.today().strftime('%Y-%m-%d'),)
        ).fetchall()
        filename = f"club_logs_today_{date.today()}.docx"
        doc_title = f"Club Logbook - {date.today()}"
    elif view_mode == 'date':
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs WHERE date = ?',
            (selected_date,)
        ).fetchall()
        filename = f"club_logs_{selected_date}.docx"
        doc_title = f"Club Logbook - {selected_date}"
    else:
        logs_data = log_db.execute(
            'SELECT user_id, date, work_done, check_in_time, check_out_time, timestamp FROM logs ORDER BY date DESC'
        ).fetchall()
        filename = "club_logs_all.docx"
        doc_title = "Club Logbook - Complete Records"

    users_data = user_db.execute('SELECT id, student_name FROM users').fetchall()
    user_map = {user['id']: user['student_name'] for user in users_data}

    # Create Word document
    doc = Document()

    # Title
    title = doc.add_heading(doc_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Generated date
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')

    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # Headers
    hdr_cells = table.rows[0].cells
    headers = ['Student Name', 'Date', 'Check-In', 'Check-Out', 'Work Summary', 'Submitted']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Add data
    for log in logs_data:
        row_cells = table.add_row().cells
        row_cells[0].text = user_map.get(log['user_id'], 'Unknown')
        row_cells[1].text = log['date']
        row_cells[2].text = log['check_in_time'] or 'N/A'
        row_cells[3].text = log['check_out_time'] or 'N/A'
        row_cells[4].text = log['work_done'] or 'No work summary'
        row_cells[5].text = log['timestamp']

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output, download_name=filename, as_attachment=True,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.errorhandler(403)
def forbidden(e):
    return render_template('403.html'), 403


if __name__ == '__main__':
    init_db()
    # For deployment, use this:
    app.run(host='0.0.0.0', port=5000)